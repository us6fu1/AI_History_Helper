"""
Экспорт тестов в формат Word (.docx) через python-docx.
Строгий чёрно-белый стиль: без цветовых акцентов, только bold/italic
для смысловых выделений.
"""
import os
import datetime
import hashlib
import random
import re
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Строки бланка «1. ___ …» — в истории часто всё склеивается пробелами, Word переносит криво.
_CHRONO_ANSWER_LINE_START = re.compile(r"(?<!\S)\d+\.\s+_{8,}")

# Разделители в заданиях на сопоставление (слева — термин, справа — пояснение).
_MATCH_PAIR_SEPARATORS = (
    " → ",
    " ⟶ ",
    " →",
    "→ ",
    "→",
    " -> ",
    "—>",
    "–>",
)

_LEFT_LABELS_MATCH = ["А", "Б", "В", "Г", "Д", "Е", "Ж", "З"]
_CHOICE_BOX = "□"


class DocxExporter:
    """Экспортирует тест в Word-документ. Только чёрно-белая палитра."""

    # Чёрный — единственный цвет текста.
    COLOR_BLACK = RGBColor(0x00, 0x00, 0x00) if DOCX_AVAILABLE else None
    # Тёмно-серый используется только для второстепенных метаданных
    # (источник, шапка для сноски).
    COLOR_GRAY = RGBColor(0x4A, 0x4A, 0x4A) if DOCX_AVAILABLE else None

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx не установлен.\n"
                "Установите: pip install python-docx"
            )

    def export(
        self,
        questions: list,
        topic: str,
        textbook_name: str,
        variant_num: int = 1,
        include_answers: bool = False,
        include_explanations: bool = False,
        include_answer_sheet: bool = True,
        output_path: Optional[str] = None,
        class_grade: str = "",
    ) -> str:
        """Создаёт .docx файл с тестом и возвращает путь к нему.

        include_answers — подсветка верных вариантов и ответы в теле теста.
        include_answer_sheet — отдельная страница «ключ» (независимо от include_answers).
        """
        doc = Document()

        self._setup_page(doc)
        self._add_header(doc, topic, textbook_name, variant_num, class_grade)
        self._add_test_instructions(doc, questions)
        self._add_separator(doc)

        for i, question in enumerate(questions, start=1):
            q = question if isinstance(question, dict) else question.to_dict()
            self._add_question(
                doc, i, q,
                include_answers=include_answers,
                include_explanations=include_explanations,
            )

        if include_answer_sheet:
            self._add_answer_sheet(doc, questions)

        if output_path is None:
            safe_topic = "".join(c for c in topic if c.isalnum() or c in " _-")[:40]
            filename = (
                f"Тест_{safe_topic}_В{variant_num}_"
                f"{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            output_path = os.path.join(os.path.expanduser("~"), "Documents", filename)

        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
        doc.save(output_path)
        return output_path


    def _setup_page(self, doc):
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.color.rgb = self.COLOR_BLACK

    def _add_header(self, doc, topic: str, textbook_name: str, variant_num: int, class_grade: str):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("ТЕСТ ПО ИСТОРИИ")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = self.COLOR_BLACK

        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(textbook_name)
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = self.COLOR_GRAY
        sub_run.italic = True

        doc.add_paragraph()

        table = doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        fields = [
            ("ФИО:", "_" * 40),
            ("Класс:", class_grade if class_grade else "_" * 25),
            ("Тема:", topic),
            ("Вариант:", str(variant_num)),
            ("Дата:", "_" * 25),
        ]

        for i, (label, value) in enumerate(fields):
            row = table.rows[i]
            label_cell = row.cells[0]
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(11)
            label_run.font.color.rgb = self.COLOR_BLACK
            label_cell.width = Cm(3)

            value_cell = row.cells[1]
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.size = Pt(11)
            value_run.font.color.rgb = self.COLOR_BLACK
            if label == "Тема:":
                value_run.bold = True

        doc.add_paragraph()

    def _add_test_instructions(self, doc, questions: list):
        q_types = {
            (q if isinstance(q, dict) else q.to_dict()).get("question_type", "test")
            for q in questions
        }

        parts = []
        if "test" in q_types:
            parts.append(
                "в заданиях с вариантами ответа поставьте галочку или крестик "
                "в квадрате рядом с тем ответом, который считаете правильным"
            )
        if "open" in q_types:
            parts.append("в открытых вопросах запишите ответ на строках под заданием")
        if "chronology" in q_types:
            parts.append("в заданиях на хронологию расположите события в правильном порядке")
        if "match" in q_types:
            parts.append(
                "в заданиях на сопоставление соедините пары или запишите соответствия"
            )
        if not parts:
            return

        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.3)
        para.paragraph_format.right_indent = Cm(0.3)
        para.paragraph_format.space_after = Pt(8)

        lead = para.add_run("Как выполнять работу: ")
        lead.bold = True
        lead.italic = True
        lead.font.size = Pt(10)
        lead.font.color.rgb = self.COLOR_GRAY

        body = para.add_run("; ".join(parts) + ".")
        body.italic = True
        body.font.size = Pt(10)
        body.font.color.rgb = self.COLOR_GRAY

    def _add_separator(self, doc):
        para = doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _split_chronology_question_text(full: str) -> tuple[list[str], list[str]]:
        """
        Возвращает абзацы формулировки и отдельно строки бланка (1. __ …), если они есть.
        """
        t = (full or "").replace("\r\n", "\n").strip()
        if not t:
            return [], []

        m = _CHRONO_ANSWER_LINE_START.search(t)
        if m:
            head = t[: m.start()].rstrip()
            tail = t[m.start() :].strip()
        else:
            mnl = re.search(r"\n\s*(\d+\.\s+_{8,})", t)
            if not mnl:
                return ([t], [])
            head = t[: mnl.start()].strip()
            tail = t[mnl.start() :].strip()

        raw_blank_lines = [ln.strip() for ln in tail.split("\n") if ln.strip()]
        if len(raw_blank_lines) <= 1:
            raw_blank_lines = [
                p.strip()
                for p in re.split(r"\s+(?=(?<!\S)\d+\.\s+_{8,})", tail)
                if p.strip()
            ]
        blanks = raw_blank_lines if raw_blank_lines else ([tail] if tail else [])

        stem_clean = head.strip()
        stem_paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", stem_clean) if p.strip()]
        if not stem_paragraphs:
            stem_paragraphs = [stem_clean] if stem_clean else []

        normalized_blanks: list[str] = []
        for b in blanks:
            mnum = re.match(r"^(\d+)\.\s*", b.strip())
            if mnum and "_" in b:
                normalized_blanks.append(f"{mnum.group(1)}. " + "_" * 46)
            else:
                normalized_blanks.append(b.strip())
        return stem_paragraphs, normalized_blanks

    @staticmethod
    def _split_match_pair(option_line: str) -> Optional[tuple[str, str]]:
        s = (option_line or "").strip()
        if not s:
            return None
        for sep in _MATCH_PAIR_SEPARATORS:
            if sep in s:
                left, right = s.split(sep, 1)
                left, right = left.strip(), right.strip()
                if left and right:
                    return left, right
        return None

    @staticmethod
    def _parse_match_options(options: list[str]) -> Optional[list[tuple[str, str]]]:
        out: list[tuple[str, str]] = []
        for opt in options:
            p = DocxExporter._split_match_pair(str(opt))
            if not p:
                return None
            out.append(p)
        return out if len(out) >= 2 else None

    def _shuffle_match_rights(self, n: int, seed: int) -> list[int]:
        rng = random.Random(seed)
        order = list(range(n))
        rng.shuffle(order)
        if n > 1:
            if all(order[i] == i for i in range(n)):
                order[0], order[1] = order[1], order[0]
        return order

    def _cell_fill(self, cell, prefix: str, body: str, *, prefix_bold: bool = True) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        if prefix:
            r0 = p.add_run(prefix)
            r0.bold = prefix_bold
            r0.font.size = Pt(11)
            r0.font.color.rgb = self.COLOR_BLACK
        r1 = p.add_run(body)
        r1.bold = False
        r1.font.size = Pt(11)
        r1.font.color.rgb = self.COLOR_BLACK

    def _add_match_table(self, doc, question_num: int, options: list[str]) -> bool:
        parsed = self._parse_match_options(options)
        if not parsed:
            return False
        n = len(parsed)
        left_texts = [p[0] for p in parsed]
        rights = [p[1] for p in parsed]
        seed_blob = hashlib.sha256(f"match:{question_num}".encode())
        seed_blob.update("\n".join(options).encode("utf-8"))
        seed = int.from_bytes(seed_blob.digest()[:8], "big")
        perm = self._shuffle_match_rights(n, seed)
        rights_shuffled = [rights[j] for j in perm]

        hint = doc.add_paragraph()
        hint.paragraph_format.left_indent = Cm(0.4)
        hint_run = hint.add_run(
            "Сопоставьте понятие в первом столбце с описанием во втором "
            "(порядок описаний перемешан). Запишите пары «буква — цифра» "
            "или соедините соответствия линиями в тетради."
        )
        hint_run.italic = True
        hint_run.font.size = Pt(10)
        hint_run.font.color.rgb = self.COLOR_GRAY

        table = doc.add_table(rows=1 + n, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_l, hdr_r = table.rows[0].cells
        self._cell_fill(hdr_l, "", "Понятие", prefix_bold=False)
        hdr_l.paragraphs[0].runs[0].bold = True
        self._cell_fill(hdr_r, "", "Описание", prefix_bold=False)
        hdr_r.paragraphs[0].runs[0].bold = True
        hdr_l.width = Cm(7.2)
        hdr_r.width = Cm(7.2)

        for i in range(n):
            row = table.rows[i + 1]
            c0, c1 = row.cells
            c0.width = Cm(7.2)
            c1.width = Cm(7.2)
            letter = _LEFT_LABELS_MATCH[i] if i < len(_LEFT_LABELS_MATCH) else str(i + 1)
            self._cell_fill(c0, f"{letter}) ", left_texts[i], prefix_bold=True)
            self._cell_fill(c1, f"{i + 1}) ", rights_shuffled[i], prefix_bold=True)

        doc.add_paragraph()
        return True
    def _add_question(self, doc, num: int, question: dict, include_answers: bool, include_explanations: bool):
        q_type = question.get("question_type", "test")
        text = question.get("text", "")
        options = question.get("options", [])
        correct = question.get("correct_answer", "")
        explanation = question.get("explanation", "")
        source_page = question.get("source_page", "")
        source_paragraph = str(question.get("source_paragraph") or "").strip()

        q_para = doc.add_paragraph()
        num_run = q_para.add_run(f"{num}. ")
        num_run.bold = True
        num_run.font.size = Pt(12)
        num_run.font.color.rgb = self.COLOR_BLACK

        chronology_blank_lines: list[str] = []
        if q_type == "chronology":
            stem_parts, chronology_blank_lines = self._split_chronology_question_text(text)
            body0 = stem_parts[0] if stem_parts else (text.strip() if text else "")
            text_run = q_para.add_run(body0)
            text_run.bold = True
            text_run.font.size = Pt(12)
            text_run.font.color.rgb = self.COLOR_BLACK
            for extra_stem in stem_parts[1:]:
                stem_para = doc.add_paragraph()
                stem_para.paragraph_format.left_indent = Cm(0.5)
                stem_run = stem_para.add_run(extra_stem)
                stem_run.bold = True
                stem_run.font.size = Pt(12)
                stem_run.font.color.rgb = self.COLOR_BLACK
        else:
            text_run = q_para.add_run(text)
            text_run.bold = True
            text_run.font.size = Pt(12)
            text_run.font.color.rgb = self.COLOR_BLACK

        src_bits: list[str] = []
        if str(source_page).strip():
            src_bits.append(str(source_page).strip())
        if source_paragraph:
            src_bits.append(source_paragraph)
        src_line = " · ".join(src_bits)

        if include_answers and src_line:
            src_para = doc.add_paragraph()
            src_run = src_para.add_run(f"    ({src_line})")
            src_run.font.size = Pt(9)
            src_run.font.color.rgb = self.COLOR_GRAY
            src_run.italic = True

        if q_type == "test" and options:
            for option in options[:4]:
                opt_para = doc.add_paragraph()
                opt_para.paragraph_format.left_indent = Cm(1)

                box_run = opt_para.add_run(f"{_CHOICE_BOX} ")
                box_run.font.size = Pt(12)
                box_run.font.color.rgb = self.COLOR_BLACK

                opt_run = opt_para.add_run(option)
                opt_run.font.size = Pt(11)
                opt_run.font.color.rgb = self.COLOR_BLACK
                
                if include_answers and correct:
                    if option.strip().lower() == correct.strip().lower():
                        opt_run.bold = True
                        opt_run.underline = True

        elif q_type == "open":
            for _ in range(3):
                line_para = doc.add_paragraph("_" * 80)
                line_para.paragraph_format.left_indent = Cm(1)
                line_para.paragraph_format.space_before = Pt(4)
                for run in line_para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = self.COLOR_GRAY

            if include_answers and correct:
                ans_para = doc.add_paragraph()
                ans_run = ans_para.add_run(f"Ответ: {correct}")
                ans_run.font.size = Pt(10)
                ans_run.font.color.rgb = self.COLOR_BLACK
                ans_run.italic = True

        elif q_type == "chronology":
            if options:
                hint = doc.add_paragraph()
                hint.paragraph_format.left_indent = Cm(1)
                hint_run = hint.add_run("Варианты (в произвольном порядке):")
                hint_run.italic = True
                hint_run.font.size = Pt(10)
                hint_run.font.color.rgb = self.COLOR_GRAY
                for opt in options:
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Cm(1)
                    opt_run = opt_para.add_run(f"• {opt}")
                    opt_run.font.size = Pt(11)
                    opt_run.font.color.rgb = self.COLOR_BLACK

            if include_answers and correct:
                ans_para = doc.add_paragraph()
                ans_run = ans_para.add_run(f"Правильный порядок/соответствие: {correct}")
                ans_run.font.size = Pt(10)
                ans_run.font.color.rgb = self.COLOR_BLACK
                ans_run.italic = True

        elif q_type == "match":
            if options and self._add_match_table(doc, num, options):
                pass
            elif options:
                for opt in options:
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Cm(1)
                    opt_run = opt_para.add_run(f"• {opt}")
                    opt_run.font.size = Pt(11)
                    opt_run.font.color.rgb = self.COLOR_BLACK

            if include_answers and correct:
                ans_para = doc.add_paragraph()
                ans_run = ans_para.add_run(f"Правильные пары: {correct}")
                ans_run.font.size = Pt(10)
                ans_run.font.color.rgb = self.COLOR_BLACK
                ans_run.italic = True

        if chronology_blank_lines:
            for ln in chronology_blank_lines:
                blank_para = doc.add_paragraph()
                blank_para.paragraph_format.keep_together = True
                blank_para.paragraph_format.left_indent = Cm(1)
                blank_para.paragraph_format.space_before = Pt(4)
                b_run = blank_para.add_run(ln)
                b_run.bold = False
                b_run.font.size = Pt(11)
                b_run.font.color.rgb = self.COLOR_BLACK

        if include_explanations and explanation:
            exp_para = doc.add_paragraph()
            exp_run = exp_para.add_run(f"Пояснение: {explanation}")
            exp_run.font.size = Pt(10)
            exp_run.font.color.rgb = self.COLOR_GRAY
            exp_run.italic = True

        doc.add_paragraph()

    def _add_answer_sheet(self, doc, questions: list):
        doc.add_page_break()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("КЛЮЧ К ТЕСТУ (ПРАВИЛЬНЫЕ ОТВЕТЫ)")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = self.COLOR_BLACK

        doc.add_paragraph()

        for i, question in enumerate(questions, start=1):
            q = question if isinstance(question, dict) else question.to_dict()
            correct = q.get("correct_answer", "")
            options = q.get("options", [])
            q_type = q.get("question_type", "test")

            if not correct:
                continue

            ans_para = doc.add_paragraph()
            num_run = ans_para.add_run(f"{i}. ")
            num_run.bold = True
            num_run.font.size = Pt(11)
            num_run.font.color.rgb = self.COLOR_BLACK

            if q_type == "test" and options:
                for opt in options[:4]:
                    if opt.strip().lower() == correct.strip().lower():
                        ans_run = ans_para.add_run(f"{_CHOICE_BOX} {correct}")
                        break
                else:
                    ans_run = ans_para.add_run(correct)
            else:
                ans_run = ans_para.add_run(correct)

            ans_run.font.size = Pt(11)
            ans_run.font.color.rgb = self.COLOR_BLACK
            ans_run.bold = True


def check_docx_available() -> bool:
    return DOCX_AVAILABLE
